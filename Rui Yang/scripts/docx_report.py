"""Word (.docx) export for the Technical and Management incident reports
(Rui Yang FYP).

Renders the SAME underlying data as report.py (technical) and
management_report.py (plain-English) into a downloadable .docx via
python-docx, so the exported file always matches what's shown on screen.

Import-only (no Streamlit page calls). Returns raw bytes so callers can feed
them straight into st.download_button without touching the filesystem.
"""
import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from report import (
    build_reasons, build_actions, attack_breakdown, top_attackers,
    per_attack_cards
)
from management_report import (
    build_overall_summary, build_incident_cards, attack_type_counts_plain
)

_ACCENT = RGBColor(0xD9, 0x77, 0x57)
_MUTED = RGBColor(0x7C, 0x7A, 0x70)

# Threat Level strings carry an emoji prefix (e.g. "🟠 High") for on-screen
# display, but Word's default document font has no colored glyph for those
# characters - it falls back to a hollow/monochrome outline instead of a
# colored circle. Rather than depend on emoji-font rendering (unreliable
# across Word versions), strip the emoji and color the word itself natively.
_LEVEL_COLOR = {
    "Critical": RGBColor(0xF0, 0x79, 0x5A),
    "High":     RGBColor(0xF0, 0xA0, 0x63),
    "Medium":   RGBColor(0xE0, 0xB6, 0x5C),
    "Low":      RGBColor(0x9F, 0xC0, 0x8A),
    "Normal":   RGBColor(0x97, 0xC0, 0xA4),
}


def _clean_level(level_text):
    """'🟠 High' -> 'High' (last whitespace-separated word)."""
    words = str(level_text).split()
    return words[-1] if words else str(level_text)


def _level_run(paragraph, level_text, bold=True):
    """Add the level word as a colored run (no emoji)."""
    word = _clean_level(level_text)
    run = paragraph.add_run(word)
    run.bold = bold
    run.font.color.rgb = _LEVEL_COLOR.get(word, _MUTED)
    return run


def _add_title(doc, title, subtitle):
    h = doc.add_heading(title, level=0)
    h.runs[0].font.color.rgb = _ACCENT
    sub = doc.add_paragraph(subtitle)
    sub.runs[0].italic = True
    sub.runs[0].font.color.rgb = _MUTED
    meta = doc.add_paragraph(
        f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    meta.runs[0].font.size = Pt(9)
    meta.runs[0].font.color.rgb = _MUTED
    doc.add_paragraph()


def _bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _to_bytes(doc):
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_technical_docx(alerts_df, overall_score, overall_level, get_ip_location):
    """Word version of report.py's Threat Analysis Report."""
    doc = Document()
    _add_title(doc, "Threat Analysis Report",
               "Hybrid IDPS — PCAP Forensic Analysis (Technical)")

    doc.add_heading("Overview", level=1)
    ov = doc.add_paragraph()
    ov.add_run(f"Threat Score: ").bold = True
    ov.add_run(f"{overall_score} / 100\n")
    ov2 = doc.add_paragraph()
    ov2.add_run("Threat Level: ").bold = True
    _level_run(ov2, overall_level)
    ov3 = doc.add_paragraph()
    ov3.add_run("Suspicious Flows: ").bold = True
    ov3.add_run(str(len(alerts_df)))

    doc.add_heading("Detected", level=1)
    _bullets(doc, [f"{atk} — {cnt} flow(s)"
                   for atk, cnt in attack_breakdown(alerts_df).most_common()])

    doc.add_heading("Possible Reasons", level=1)
    _bullets(doc, build_reasons(alerts_df))

    doc.add_heading("Suggested Actions", level=1)
    _bullets(doc, build_actions(alerts_df))

    doc.add_heading("Top Attacking Sources", level=1)
    attackers = top_attackers(alerts_df, get_ip_location, limit=5)
    if attackers:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Source IP", "Flows", "Origin"
        for a in attackers:
            row = table.add_row().cells
            row[0].text = str(a.get("ip", "-"))
            row[1].text = str(a.get("flows", "-"))
            row[2].text = str(a.get("origin", "-"))
    else:
        doc.add_paragraph("No attacking sources to list.")

    # No 5-item cap here, unlike the on-screen view - that limit exists only
    # because of live-screen space; a downloaded document has no such
    # constraint and should stand alone with full detail (it also can't tell
    # the reader to "see the app," since by the time they're reading this
    # file there may be no app open at all).
    doc.add_heading("Per-Attack Breakdown", level=1)
    for card in per_attack_cards(alerts_df, get_ip_location):
        h = doc.add_paragraph(style="Heading 2")
        h.add_run(f"{card['reason_name']} — Score {card['score']} (")
        _level_run(h, card['level'])
        h.add_run(")")
        doc.add_paragraph(
            f"{card['src']} → {card['dst']} : port {card['port']}"
        )
        p = doc.add_paragraph()
        p.add_run("Why: ").bold = True
        p.add_run(str(card['why']))
        p2 = doc.add_paragraph()
        p2.add_run("Action: ").bold = True
        p2.add_run(str(card['action']))
        origin = doc.add_paragraph(f"Origin: {card['origin']}")
        origin.runs[0].italic = True

    return _to_bytes(doc)


def build_management_docx(alerts_df, overall_score, overall_level, get_ip_location):
    """Word version of management_report.py's plain-English incident report."""
    doc = Document()
    _add_title(doc, "Incident Summary",
               "Hybrid IDPS — Non-Technical Briefing (Management)")

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(
        build_overall_summary(alerts_df, overall_score, _clean_level(overall_level))
    )

    doc.add_heading("What Was Detected", level=1)
    _bullets(doc, [f"{label} — {cnt} event" if cnt == 1 else f"{label} — {cnt} events"
                   for label, cnt in attack_type_counts_plain(alerts_df).most_common()])

    doc.add_heading("Incident Details", level=1)
    for card in build_incident_cards(alerts_df, get_ip_location):
        doc.add_heading(f"Source: {card['origin']}", level=2)
        meta = doc.add_paragraph()
        meta.add_run("When: ").bold = True
        meta.add_run(f"{card['start']} to {card['end']}    ")
        meta.add_run("Rating: ").bold = True
        _level_run(meta, card['level'])

        for label, key in (
            ("What happened", "what"),
            ("Data exposure", "exposure"),
            ("Severity", "severity_sentence"),
            ("Recommended next step / lesson learnt", "takeaway"),
        ):
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(str(card[key]))

        if card.get("prior"):
            _prior_word = "prior incident" if card['prior'] == 1 else "prior incidents"
            note = doc.add_paragraph(
                f"This source has {card['prior']} {_prior_word} on record."
            )
            note.runs[0].italic = True
            note.runs[0].font.color.rgb = _MUTED
        doc.add_paragraph()

    return _to_bytes(doc)
